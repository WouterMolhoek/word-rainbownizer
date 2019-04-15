from docx import Document
from docx.shared import RGBColor

doc = Document('test-file.docx')

# All the paragraphs stored in an array
allTxt = doc.paragraphs


# Function to delete a paragraph
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


totalWords = []
newWordList = []

# Red, Orange, Yellow, Green, Blue, Violet
colors = [RGBColor(255, 0, 0), RGBColor(255, 117, 0), RGBColor(255, 249, 0), RGBColor(76, 223, 0), RGBColor(26, 103, 236), RGBColor(151, 23, 238)]

# Loop through the existing paragraphs, delete the old ones and replace them with the colorful ones.
for p in allTxt:
    # Remove existing paragraphs
    delete_paragraph(p)
    # Split every word between a space
    words = p.text.split(' ')

    # Loop through the words an push them in the totalWords array
    for word in words:
        totalWords.append(word)

part = int(len(totalWords) / 6)


def add_words(begin, end, color):
    for i in range(begin, end):
        newWordList.append(totalWords[i])

    run = doc.add_paragraph().add_run(newWordList)
    font = run.font

    # Set font color to red
    font.color.rgb = color
    # Clear the previous array
    newWordList.clear()


add_words(0, part, colors[0])
add_words(part, part * 2, colors[1])
add_words(part * 2, part * 3, colors[2])
add_words(part * 3, part * 4, colors[3])
add_words(part * 4, part * 5, colors[4])
add_words(part * 5, part * 6, colors[5])

# Create a new file and save it
doc.save('rainbow.docx')
